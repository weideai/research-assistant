import csv
import io
import json
import os
import re
import tempfile
import zipfile
from datetime import date, datetime
from pathlib import Path
from xml.sax.saxutils import escape


REPORT_TEMPLATES = {
    "research": {
        "label": "科研档案",
        "kicker": "RESEARCH RECORD",
        "accent": "2166F3",
        "text_accent": "174EA6",
        "soft": "EDF3FF",
        "include_hash": True,
    },
    "notebook": {
        "label": "实验记录本",
        "kicker": "LAB NOTEBOOK",
        "accent": "167C80",
        "text_accent": "0F6265",
        "soft": "E9F7F4",
        "include_hash": False,
    },
    "compact": {
        "label": "简洁结果报告",
        "kicker": "RESULT SUMMARY",
        "accent": "A15C00",
        "text_accent": "7A4300",
        "soft": "FFF4D6",
        "include_hash": False,
    },
}

# NOTE: `include_hash` is reserved for the structured report descriptor (P1). Reports
# intentionally omit per-file names and hashes — they show the folder name plus a file
# count and route the reader to the file centre instead. See docs/UX-IA-V3.md §3.

MAX_REPORT_ATTACHMENT_PREVIEWS = 2


def report_template_choices():
    return tuple((key, value["label"]) for key, value in REPORT_TEMPLATES.items())


def resolve_report_template(template_key):
    """Return the report profile for a key, falling back to the default template."""
    key = (template_key or "").strip().lower()
    return REPORT_TEMPLATES.get(key, REPORT_TEMPLATES["research"])


def _date_value(value):
    return value.isoformat() if value else None


def _text(value, fallback=""):
    if value is None:
        return fallback
    return str(value)


def _active(items):
    return [item for item in items if not getattr(item, "is_deleted", False)]


def _record_sort_key(record):
    return record.record_date or date.min, record.id or 0


def _batch_sort_key(batch):
    return batch.start_date or date.min, batch.id or 0


def _step_sort_key(step):
    return step.position or 0, step.id or 0


def _plan_step_payload(step):
    return {
        "id": step.id,
        "position": step.position,
        "title": step.title,
        "description": step.description,
        "operator": step.operator,
        "planned_date": _date_value(step.planned_date),
    }


def _batch_step_payload(step):
    return {
        "id": step.id,
        "source_step_id": step.source_step_id,
        "position": step.position,
        "title": step.title,
        "description": step.description,
        "operator": step.operator,
        "planned_date": _date_value(step.planned_date),
        "completed_date": _date_value(step.completed_date),
        "is_done": step.is_done,
    }


def _parameter_payload(parameter):
    return {
        "position": parameter.position,
        "name": parameter.name,
        "value": parameter.value,
        "unit": parameter.unit,
        "notes": parameter.notes,
    }


def _sample_usage_payload(usage):
    return {
        "sample_code": usage.sample.sample_code,
        "sample_type": usage.sample.sample_type,
        "source": usage.sample.source,
        "location": usage.sample.location,
        "quantity": usage.sample.quantity,
        "status": usage.sample.status,
        "role": usage.role,
        "amount_used": usage.amount_used,
        "notes": usage.notes,
    }


def _attachment_payload(attachment):
    return {
        "id": attachment.id,
        "category": attachment.category,
        "original_name": attachment.original_name,
        "relative_path": attachment.relative_path,
        "version_number": attachment.version_number,
        "size_bytes": attachment.size_bytes,
        "mime_type": attachment.mime_type,
        "is_previewable_image": attachment.is_previewable_image,
        "sha256": attachment.sha256,
        "tags": attachment.tags,
        "description": attachment.description,
        "storage_mode": attachment.storage_mode,
        "link_status": attachment.link_status,
    }


def _record_attachment_folder_name(record_id):
    return f"record-{record_id}"


def _attachment_folder_note(record_id, attachments):
    count = len(attachments)
    folder_name = _record_attachment_folder_name(record_id)
    if not count:
        return folder_name, "暂无原始数据或附件。"
    return (
        folder_name,
        f"共 {count} 个文件。报告仅展示最多 {MAX_REPORT_ATTACHMENT_PREVIEWS} 张核心结果缩略图，"
        "其余原始数据请在文件中心打开此文件夹查看。",
    )


def _resolved_attachment_previews(attachments, attachment_lookup, attachment_path_resolver):
    if not attachment_path_resolver:
        return []
    previews = []
    for attachment in attachments:
        if len(previews) >= MAX_REPORT_ATTACHMENT_PREVIEWS:
            break
        if not attachment.get("is_previewable_image"):
            continue
        attachment_item = attachment_lookup.get(attachment.get("id"))
        if attachment_item is None:
            continue
        try:
            source = Path(attachment_path_resolver(attachment_item))
        except Exception:
            continue
        if source.is_file():
            previews.append((attachment, source))
    return previews


def _record_payload(record, batch):
    return {
        "id": record.id,
        "batch_id": batch.id if batch else None,
        "batch_code": (batch.batch_code or f"EXEC-{batch.id}") if batch else "HISTORY-UNASSIGNED",
        "record_date": _date_value(record.record_date),
        "operator": record.operator,
        "conditions": record.conditions,
        "content": record.content,
        "result": record.result,
        "remark": record.remark,
        "lifecycle_status": record.lifecycle_status,
        "finalized_at": _date_value(record.finalized_at),
        "parameters": [_parameter_payload(parameter) for parameter in record.parameters],
        "attachments": [
            _attachment_payload(attachment)
            for attachment in sorted(
                _active(record.attachments),
                key=lambda attachment: (attachment.relative_path.lower(), attachment.version_number),
            )
        ],
    }


def execution_groups(payload):
    """Return real executions plus one explicit compatibility group for orphaned records."""
    groups = list(payload["batches"])
    if payload["unassigned_records"]:
        groups.append({
            "id": None,
            "batch_code": "HISTORY-UNASSIGNED",
            "repeat_kind": "历史数据",
            "repeat_number": None,
            "group_name": "待归档",
            "operator": "",
            "status": "待修复",
            "start_date": None,
            "end_date": None,
            "summary": "这些旧记录缺少有效的实验批次归属；导出保留数据，但建议先运行数据库升级完成归档。",
            "conclusion": "",
            "requires_repeat": False,
            "steps": [],
            "actual_parameters": [],
            "sample_usages": [],
            "records": payload["unassigned_records"],
            "is_unassigned": True,
        })
    return groups


def experiment_payload(item, exported_at=None):
    """Return one stable data shape shared by every structured export format."""
    exported_at = exported_at or datetime.now()
    batches = sorted(_active(item.batches), key=_batch_sort_key)
    batch_by_id = {batch.id: batch for batch in batches}
    records = sorted(_active(item.records), key=_record_sort_key)
    record_payloads = {
        record.id: _record_payload(record, batch_by_id.get(record.batch_id))
        for record in records
    }
    assigned_record_ids = set()
    batch_payloads = []
    for batch in batches:
        batch_records = [
            record_payloads[record.id]
            for record in records
            if record.batch_id == batch.id
        ]
        assigned_record_ids.update(record["id"] for record in batch_records)
        batch_payloads.append({
            "id": batch.id,
            "batch_code": batch.batch_code or f"EXEC-{batch.id}",
            "repeat_kind": batch.repeat_kind,
            "repeat_number": batch.repeat_number,
            "group_name": batch.group_name,
            "operator": batch.operator,
            "status": batch.status,
            "start_date": _date_value(batch.start_date),
            "end_date": _date_value(batch.end_date),
            "summary": batch.summary,
            "conclusion": batch.conclusion,
            "requires_repeat": batch.requires_repeat,
            "steps": [
                _batch_step_payload(step)
                for step in sorted(batch.steps, key=_step_sort_key)
            ],
            "actual_parameters": [_parameter_payload(parameter) for parameter in batch.actual_parameters],
            "sample_usages": [_sample_usage_payload(usage) for usage in batch.sample_usages],
            "records": batch_records,
            "is_unassigned": False,
        })
    unassigned_records = [
        record_payloads[record.id]
        for record in records
        if record.id not in assigned_record_ids
    ]
    return {
        "schema_version": 3,
        "exported_at": exported_at.isoformat(timespec="seconds"),
        "experiment": {
            "id": item.id,
            "project_id": item.project_id,
            "project_title": item.project.title if item.project else "",
            "title": item.title,
            "code": item.code,
            "status": item.status,
            "owner": item.owner,
            "start_date": _date_value(item.start_date),
            "end_date": _date_value(item.end_date),
            "objective": item.objective,
        },
        "samples": [_sample_usage_payload(usage) for usage in item.sample_usages],
        "plan_parameters": [_parameter_payload(parameter) for parameter in item.plan_parameters],
        "steps": [
            _plan_step_payload(step)
            for step in sorted(item.steps, key=_step_sort_key)
        ],
        "batches": batch_payloads,
        "unassigned_records": unassigned_records,
        "records": [record_payloads[record.id] for record in records],
    }


def build_json_export(item):
    return json.dumps(experiment_payload(item), ensure_ascii=False, indent=2).encode("utf-8")


def _markdown_value(value, fallback="未填写"):
    return str(value).strip() if value not in (None, "") else fallback


def _markdown_cell(value, fallback="—"):
    text = _markdown_value(value, fallback).replace("|", "\\|")
    return "<br>".join(part.strip() for part in text.splitlines())


def _markdown_table(headers, rows, empty="暂无数据。"):
    if not rows:
        return [empty]
    lines = [
        "| " + " | ".join(headers) + " |",
        "| " + " | ".join("---" for _ in headers) + " |",
    ]
    lines.extend("| " + " | ".join(_markdown_cell(value) for value in row) + " |" for row in rows)
    return lines


def _size_label(size_bytes):
    size = int(size_bytes or 0)
    if size < 1024:
        return f"{size} B"
    if size < 1024 * 1024:
        return f"{size / 1024:.1f} KB"
    if size < 1024 * 1024 * 1024:
        return f"{size / (1024 * 1024):.1f} MB"
    return f"{size / (1024 * 1024 * 1024):.1f} GB"


def _markdown_from_payload(payload):
    experiment = payload["experiment"]
    lines = [
        f"# {experiment['title']}",
        "",
        "> 实验导出报告 · R/LAB Research Assistant",
        "> 数据按当前实验快照生成，原始附件请优先使用 ZIP 完整归档保存。",
        "",
        "## 目录",
        "",
        "1. [实验概览](#实验概览)",
        "2. [实验目的](#实验目的)",
        "3. [关联样本与计划参数](#关联样本与计划参数)",
        "4. [计划步骤定义](#计划步骤定义)",
        "5. [实验批次与过程记录](#实验批次与过程记录)",
        "",
        "---",
        "",
        "## 实验概览",
        "",
    ]
    lines.extend(_markdown_table(("字段", "内容"), [
        ("实验编号", _markdown_value(experiment["code"], "未设置")),
        ("状态", experiment["status"]),
        ("负责人", _markdown_value(experiment["owner"])),
        ("计划开始", _markdown_value(experiment["start_date"], "未安排")),
        ("计划结束", _markdown_value(experiment["end_date"], "未安排")),
        ("导出时间", payload["exported_at"].replace("T", " ")),
    ]))
    lines.extend([
        "", "## 实验目的", "",
        "> " + _markdown_value(experiment["objective"]).replace("\n", "\n> "), "",
        "## 关联样本与计划参数", "", "### 关联样本", "",
    ])
    lines.extend(_markdown_table(("样本编号", "类型", "用途", "使用量", "位置", "状态", "备注"), [
        (sample["sample_code"], sample["sample_type"], sample["role"], sample["amount_used"],
         sample["location"], sample["status"], sample["notes"])
        for sample in payload["samples"]
    ], "暂无关联样本。"))
    lines.extend(["", "### 计划参数", ""])
    lines.extend(_markdown_table(("序号", "参数", "数值", "单位", "说明"), [
        (parameter["position"], parameter["name"], parameter["value"], parameter["unit"], parameter["notes"])
        for parameter in payload["plan_parameters"]
    ], "暂无结构化计划参数。"))
    lines.extend(["", "## 计划步骤定义", ""])
    lines.extend(_markdown_table(("序号", "步骤", "计划执行人", "计划日期", "说明"), [
        (step["position"], step["title"], step["operator"],
         step["planned_date"] or "未安排", step["description"])
        for step in payload["steps"]
    ], "暂无计划步骤。"))
    lines.extend(["", "## 实验批次与过程记录", ""])

    groups = execution_groups(payload)
    if not groups:
        lines.append("暂无实验批次。")
    for execution_index, execution in enumerate(groups, start=1):
        lines.extend([
            f"### 批次 {execution_index:02d} · {execution['batch_code']}",
            "",
        ])
        if execution.get("is_unassigned"):
            lines.extend(["> 注意：以下为历史未归档记录。数据已保留，请运行数据库升级完成批次归属修复。", ""])
        lines.extend(_markdown_table(("字段", "内容"), [
            ("批次编号", execution["batch_code"]),
            ("重复类型", execution["repeat_kind"]),
            ("重复序号", execution["repeat_number"]),
            ("实验分组", execution["group_name"]),
            ("实验人员", execution["operator"]),
            ("状态", execution["status"]),
            ("实际开始", execution["start_date"]),
            ("实际结束", execution["end_date"]),
            ("建议重复", "是" if execution["requires_repeat"] else "否"),
        ]))
        lines.extend(["", "#### 批次摘要", "", _markdown_value(execution["summary"]), ""])
        if execution["conclusion"]:
            lines.extend(["#### 批次结论", "", execution["conclusion"], ""])
        lines.extend(["#### 批次步骤", ""])
        lines.extend(_markdown_table(
            ("序号", "状态", "步骤", "执行人", "计划日期", "完成日期", "说明"),
            [
                (
                    step["position"], "已完成" if step["is_done"] else "待完成",
                    step["title"], step["operator"], step["planned_date"] or "未安排",
                    step["completed_date"] or "未完成", step["description"],
                )
                for step in execution["steps"]
            ],
            "暂无批次步骤。",
        ))
        lines.append("")
        if execution["actual_parameters"]:
            lines.extend(["#### 实际参数", ""])
            lines.extend(_markdown_table(("参数", "数值", "单位", "说明"), [
                (parameter["name"], parameter["value"], parameter["unit"], parameter["notes"])
                for parameter in execution["actual_parameters"]
            ]))
            lines.append("")

        if not execution["records"]:
            lines.extend(["暂无过程记录。", ""])
        for record_index, record in enumerate(execution["records"], start=1):
            lines.extend([
                f"#### 过程记录 {execution_index:02d}.{record_index:02d} · {record['record_date']} · {record['result']}",
                "",
                f"**实验人员：** {_markdown_value(record['operator'])}",
                "",
                "##### 结构化参数",
                "",
            ])
            lines.extend(_markdown_table(("参数", "数值", "单位", "说明"), [
                (parameter["name"], parameter["value"], parameter["unit"], parameter["notes"])
                for parameter in record["parameters"]
            ], "暂无结构化过程记录参数。"))
            if record["parameters"]:
                parameter_summary = "；".join(
                    f"{parameter['name']}：{parameter['value']}"
                    f"{' ' + parameter['unit'] if parameter['unit'] else ''}"
                    for parameter in record["parameters"]
                )
                lines.extend(["", f"**参数摘要：** {parameter_summary}"])
            lines.extend([
                "", "##### 实验条件", "",
                "> " + _markdown_value(record["conditions"]).replace("\n", "\n> "),
                "", "##### 实验过程", "", record["content"] or "未填写。",
                "", "##### 结论与备注", "", "> " + _markdown_value(record["remark"]), "",
            ])
            if record["attachments"]:
                lines.extend(["##### 结果与数据文件", ""])
                lines.extend(_markdown_table(
                    ("分类", "文件夹 / 文件", "版本", "大小", "SHA-256", "标签", "说明"),
                    [
                        (attachment["category"], f"`{attachment['relative_path']}`",
                         f"v{attachment['version_number']}", _size_label(attachment["size_bytes"]),
                         attachment["sha256"] or "旧文件未计算", attachment["tags"], attachment["description"])
                        for attachment in record["attachments"]
                    ],
                ))
                lines.append("")
        lines.extend(["---", ""])

    lines.extend([
        "*导出完成。请结合实验室 SOP 对剂量、统计结论和临床相关内容进行人工核验。*"
    ])
    return "\ufeff" + "\n".join(lines).rstrip() + "\n"


def build_markdown_export(item):
    return _markdown_from_payload(experiment_payload(item))


def _set_docx_font(document, profile=None):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Mm, Pt, RGBColor

    profile = profile or REPORT_TEMPLATES["research"]
    text_accent = profile["text_accent"]
    style_specs = {
        "Normal": (10.5, "17212B", False),
        "Title": (26, "0F172A", True),
        "Heading 1": (17, "102A43", True),
        "Heading 2": (13.5, text_accent, True),
        "Heading 3": (11.5, "243B53", True),
    }
    for style_name, (font_size, color, bold) in style_specs.items():
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(font_size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = bold
        style._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.paragraph_format.space_after = Pt(7)
        style.paragraph_format.line_spacing = 1.25
    document.styles["Heading 1"].paragraph_format.space_before = Pt(18)
    document.styles["Heading 1"].paragraph_format.keep_with_next = True
    document.styles["Heading 2"].paragraph_format.space_before = Pt(13)
    document.styles["Heading 2"].paragraph_format.keep_with_next = True

    section = document.sections[0]
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(17)
    section.left_margin = Mm(19)
    section.right_margin = Mm(19)
    header = section.header.paragraphs[0]
    header.text = f"R/LAB  ·  {profile['kicker']}"
    header.style = document.styles["Normal"]
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.bold = True
    header.runs[0].font.color.rgb = RGBColor.from_string(text_accent)
    footer = section.footer.paragraphs[0]
    footer.alignment = 2
    footer.add_run("实验报告  ·  ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)
    for run in footer.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor.from_string("52616B")


def _shade_docx_cell(cell, fill):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    properties = cell._tc.get_or_add_tcPr()
    shading = properties.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        properties.append(shading)
    shading.set(qn("w:fill"), fill)


def _format_docx_table(table, profile=None):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    profile = profile or REPORT_TEMPLATES["research"]
    table.autofit = True
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header_properties.append(repeat)
    for cell in table.rows[0].cells:
        _shade_docx_cell(cell, profile["accent"])
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            run.font.size = Pt(9.5)
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
            run._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    for row_index, row in enumerate(table.rows[1:], start=1):
        if row_index % 2 == 0:
            for cell in row.cells:
                _shade_docx_cell(cell, "F3F6F7")
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor.from_string("17212B")


def _add_docx_callout(profile, document, label, text, fill=None):
    from docx.shared import Pt, RGBColor

    profile = profile or REPORT_TEMPLATES["research"]
    table = document.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.cell(0, 0)
    _shade_docx_cell(cell, fill or profile["soft"])
    heading = cell.paragraphs[0]
    label_run = heading.add_run(label)
    label_run.bold = True
    label_run.font.color.rgb = RGBColor.from_string(profile["text_accent"])
    body = cell.add_paragraph(_text(text, "未填写。") or "未填写。")
    body.paragraph_format.space_after = Pt(5)
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def _add_docx_table(profile, document, headers, rows):
    if not rows:
        document.add_paragraph("暂无数据。")
        return
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = _text(value, "-") or "-"
    _format_docx_table(table, profile)
    return table


def _add_docx_attachment_section(
    profile, document, record_id, attachments, attachment_lookup,
    attachment_path_resolver=None,
):
    from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.image.image import Image as DocxImage
    from docx.shared import Emu, Mm, Pt, RGBColor

    previews = _resolved_attachment_previews(
        attachments, attachment_lookup, attachment_path_resolver,
    )
    if previews:
        preview_label = document.add_paragraph()
        preview_run = preview_label.add_run("核心结果预览")
        preview_run.bold = True
        preview_run.font.color.rgb = RGBColor.from_string("243B53")
        preview_label.paragraph_format.space_after = Pt(5)

        preview_table = document.add_table(rows=1, cols=len(previews))
        preview_table.alignment = WD_TABLE_ALIGNMENT.LEFT
        preview_table.autofit = False
        max_width = Mm(72 if len(previews) > 1 else 100)
        max_height = Mm(62)
        for cell, (_, source) in zip(preview_table.rows[0].cells, previews):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _shade_docx_cell(cell, "F7F9FA")
            picture_paragraph = cell.paragraphs[0]
            picture_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                image_bytes = source.read_bytes()
                image = DocxImage.from_file(io.BytesIO(image_bytes))
                scale = min(max_width / image.width, max_height / image.height, 1)
                picture_paragraph.add_run().add_picture(
                    io.BytesIO(image_bytes),
                    width=Emu(int(image.width * scale)),
                    height=Emu(int(image.height * scale)),
                )
            except Exception:
                picture_paragraph.add_run("预览暂不可用")
        document.add_paragraph().paragraph_format.space_after = Pt(0)

    folder_name, folder_note = _attachment_folder_note(record_id, attachments)
    _add_docx_callout(
        profile,
        document,
        "完整实验文件夹",
        f"{folder_name}  ·  {folder_note}",
        "F3F6F7",
    )


def build_docx_export(item, attachment_path_resolver=None, template_key="research"):
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    profile = resolve_report_template(template_key)
    payload = experiment_payload(item)
    experiment = payload["experiment"]
    attachment_lookup = {
        attachment.id: attachment
        for record in _active(item.records)
        for attachment in _active(record.attachments)
    }
    document = Document()
    _set_docx_font(document, profile)
    document.core_properties.title = experiment["title"]
    document.core_properties.subject = "实验计划、实验批次与过程记录导出报告"
    document.core_properties.author = "R/LAB Research Assistant"
    title = document.add_heading(experiment["title"], 0)
    title.paragraph_format.space_after = Pt(4)
    subtitle = document.add_paragraph("EXPERIMENT REPORT  /  实验计划、批次与过程记录")
    subtitle.runs[0].font.size = Pt(9)
    subtitle.runs[0].font.bold = True
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(profile["text_accent"])
    subtitle.paragraph_format.space_after = Pt(15)

    document.add_heading("01  实验概览", level=1)
    _add_docx_table(profile, document, ("字段", "内容"), (
        ("实验编号", experiment["code"] or "未设置"),
        ("状态", experiment["status"]),
        ("负责人", experiment["owner"]),
        ("计划开始", experiment["start_date"]),
        ("计划结束", experiment["end_date"]),
        ("导出时间", payload["exported_at"].replace("T", " ")),
    ))
    document.add_heading("02  实验目的", level=1)
    _add_docx_callout(profile, document, "OBJECTIVE  /  研究目的", experiment["objective"])

    document.add_heading("03  样本与计划参数", level=1)
    document.add_heading("3.1  关联样本", level=2)
    _add_docx_table(profile, document, ("样本编号", "类型", "用途", "使用量", "位置", "状态", "备注"), [
        (sample["sample_code"], sample["sample_type"], sample["role"], sample["amount_used"],
         sample["location"], sample["status"], sample["notes"])
        for sample in payload["samples"]
    ])

    document.add_heading("3.2  计划参数", level=2)
    _add_docx_table(profile, document, ("序号", "参数", "数值", "单位", "说明"), [
        (parameter["position"], parameter["name"], parameter["value"], parameter["unit"], parameter["notes"])
        for parameter in payload["plan_parameters"]
    ])

    document.add_heading("04  计划步骤定义", level=1)
    _add_docx_table(profile, document, ("序号", "步骤", "计划执行人", "计划日期", "说明"), [
        (step["position"], step["title"], step["operator"],
         step["planned_date"], step["description"])
        for step in payload["steps"]
    ])

    document.add_page_break()
    document.add_heading("05  实验批次与过程记录", level=1)
    groups = execution_groups(payload)
    if not groups:
        document.add_paragraph("暂无实验批次。")
    for execution_index, execution in enumerate(groups, start=1):
        if execution_index > 1:
            document.add_page_break()
        document.add_heading(
            f"批次 {execution_index:02d}  /  {execution['batch_code']}", level=2
        )
        if execution.get("is_unassigned"):
            _add_docx_callout(profile, document,
                "DATA INTEGRITY  /  历史未归档",
                "以下记录缺少有效的实验批次归属。数据已保留，请运行数据库升级完成修复。",
                "FFF4D6",
            )
        _add_docx_table(profile, document, ("字段", "内容"), (
            ("批次编号", execution["batch_code"]),
            ("重复类型", execution["repeat_kind"]),
            ("重复序号", execution["repeat_number"]),
            ("实验分组", execution["group_name"]),
            ("实验人员", execution["operator"]),
            ("状态", execution["status"]),
            ("实际开始", execution["start_date"]),
            ("实际结束", execution["end_date"]),
            ("建议重复", "是" if execution["requires_repeat"] else "否"),
        ))
        _add_docx_callout(profile, document, "SUMMARY  /  批次摘要", execution["summary"], "F3F6F7")
        if execution["conclusion"]:
            _add_docx_callout(profile, document, "CONCLUSION  /  批次结论", execution["conclusion"], "FFF4D6")
        document.add_heading(f"5.{execution_index}.1  批次步骤", level=3)
        _add_docx_table(profile, document, ("序号", "状态", "步骤", "执行人", "计划日期", "完成日期", "说明"), [
            (
                step["position"], "已完成" if step["is_done"] else "待完成",
                step["title"], step["operator"], step["planned_date"],
                step["completed_date"], step["description"],
            )
            for step in execution["steps"]
        ])
        if execution["actual_parameters"]:
            document.add_heading(f"5.{execution_index}.2  实际参数", level=3)
            _add_docx_table(profile, document, ("参数", "数值", "单位", "说明"), [
                (parameter["name"], parameter["value"], parameter["unit"], parameter["notes"])
                for parameter in execution["actual_parameters"]
            ])
        if not execution["records"]:
            document.add_paragraph("暂无过程记录。")
        for record_index, record in enumerate(execution["records"], start=1):
            document.add_heading(
                f"过程记录 {execution_index:02d}.{record_index:02d}  /  {record['record_date']}",
                level=3,
            )
            summary = document.add_paragraph()
            summary.alignment = WD_ALIGN_PARAGRAPH.LEFT
            summary.add_run(
                f"结果：{record['result']}  ·  实验人员：{record['operator'] or '未填写'}"
            ).bold = True
            document.add_paragraph("结构化参数").runs[0].bold = True
            _add_docx_table(profile, document, ("参数", "数值", "单位", "说明"), [
                (parameter["name"], parameter["value"], parameter["unit"], parameter["notes"])
                for parameter in record["parameters"]
            ])
            document.add_paragraph("实验条件").runs[0].bold = True
            _add_docx_callout(profile, document, "CONDITIONS  /  实验条件", record["conditions"], "F3F6F7")
            document.add_paragraph("实验过程").runs[0].bold = True
            document.add_paragraph(record["content"] or "未填写。")
            document.add_paragraph("结论与备注").runs[0].bold = True
            _add_docx_callout(profile, document, "CONCLUSION  /  结论与后续", record["remark"], "FFF4D6")
            document.add_paragraph("结果与数据文件").runs[0].bold = True
            _add_docx_attachment_section(
                profile, document, record["id"], record["attachments"], attachment_lookup,
                attachment_path_resolver,
            )

    document.add_paragraph()
    review = document.add_paragraph("人工核验提示：剂量、统计结论和临床相关解释须结合实验室 SOP 与原始数据复核。")
    review.runs[0].italic = True
    review.runs[0].font.size = Pt(8.5)
    review.runs[0].font.color.rgb = RGBColor.from_string("52616B")

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


PDF_FONT_CANDIDATES = (
    r"C:\Windows\Fonts\msyh.ttc",
    r"C:\Windows\Fonts\simhei.ttf",
    r"C:\Windows\Fonts\Deng.ttf",
    r"C:\Windows\Fonts\NotoSansSC-VF.ttf",
    r"/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    r"/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
    r"/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
    r"/usr/share/fonts/truetype/arphic/uming.ttc",
    r"/System/Library/Fonts/PingFang.ttc",
    r"/System/Library/Fonts/Hiragino Sans GB.ttc",
)

PDF_FONT_MISSING_MESSAGE = (
    "PDF 导出需要一个包含中文字形的字体，但系统中没有找到。"
    "请安装 Noto Sans CJK（Debian/Ubuntu：apt install fonts-noto-cjk），"
    "或用环境变量 RESEARCH_ASSISTANT_PDF_FONT 指定一个 .ttf/.ttc 字体文件路径。"
    "导出 Word 或 Markdown 不受影响。"
)


class PdfFontMissingError(RuntimeError):
    """Raised when no CJK-capable font is available for PDF export."""


def find_pdf_font_path():
    """Return the first usable CJK font path, or None when none is installed."""
    configured = os.getenv("RESEARCH_ASSISTANT_PDF_FONT", "").strip()
    for candidate in (configured, *PDF_FONT_CANDIDATES):
        if candidate and Path(candidate).is_file():
            return candidate
    return None


def _pdf_font_name():
    """Register and return a CJK-capable font, or fail loudly.

    Falling back to a Latin-only face such as Helvetica silently drops every CJK
    glyph, so the user receives a blank-looking PDF with no indication that the
    export was unusable. Refusing the export is the honest outcome.
    """
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_name = "ResearchAssistantCJK"
    if font_name in pdfmetrics.getRegisteredFontNames():
        return font_name
    configured = os.getenv("RESEARCH_ASSISTANT_PDF_FONT", "").strip()
    failures = []
    for candidate in (configured, *PDF_FONT_CANDIDATES):
        if not candidate:
            continue
        path = Path(candidate)
        if not path.is_file():
            continue
        try:
            pdfmetrics.registerFont(TTFont(font_name, str(path)))
            return font_name
        except Exception as exc:
            failures.append(f"{candidate}: {exc}")
            continue
    detail = ("  尝试过的字体：" + "；".join(failures)) if failures else ""
    raise PdfFontMissingError(PDF_FONT_MISSING_MESSAGE + detail)


def _pdf_value(value, fallback="未填写"):
    text = _text(value, fallback).strip()
    if not text:
        text = fallback
    return escape(text).replace("\r\n", "\n").replace("\n", "<br/>")


def build_pdf_export(item, attachment_path_resolver, template_key="research"):
    """Render a readable PDF from the shared experiment export payload."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.lib.utils import ImageReader
        from reportlab.platypus import (
            Image, KeepTogether, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle,
        )
    except ImportError as exc:
        raise RuntimeError("PDF 导出需要安装 reportlab 依赖。") from exc

    profile = REPORT_TEMPLATES.get(template_key, REPORT_TEMPLATES["research"])
    payload = experiment_payload(item)
    experiment = payload["experiment"]
    attachment_lookup = {
        attachment.id: attachment
        for record in _active(item.records)
        for attachment in _active(record.attachments)
    }
    font_name = _pdf_font_name()
    accent = colors.HexColor(f"#{profile['accent']}")
    accent_text = colors.HexColor(f"#{profile['text_accent']}")
    soft = colors.HexColor(f"#{profile['soft']}")
    ink = colors.HexColor("#17212B")
    muted = colors.HexColor("#475467")
    line = colors.HexColor("#CBD5DC")
    output = io.BytesIO()
    document = SimpleDocTemplate(
        output,
        pagesize=A4,
        rightMargin=16 * mm,
        leftMargin=16 * mm,
        topMargin=18 * mm,
        bottomMargin=17 * mm,
        title=_text(experiment["title"], "实验记录"),
        author="R/LAB Research Assistant",
    )

    body = ParagraphStyle(
        "ResearchBody", fontName=font_name, fontSize=9.3, leading=14,
        textColor=ink, spaceAfter=5,
    )
    small = ParagraphStyle(
        "ResearchSmall", parent=body, fontSize=7.8, leading=11, textColor=muted,
    )
    kicker = ParagraphStyle(
        "ResearchKicker", parent=small, fontSize=8.2, leading=10.5, textColor=accent_text,
        tracking=1.4, spaceAfter=9,
    )
    cover = ParagraphStyle(
        "ResearchCover", parent=body, fontSize=24, leading=31, textColor=ink,
        alignment=TA_LEFT, spaceAfter=12,
    )
    title = ParagraphStyle(
        "ResearchTitle", parent=body, fontSize=15, leading=20, textColor=ink,
        spaceBefore=8, spaceAfter=8,
    )
    section = ParagraphStyle(
        "ResearchSection", parent=body, fontSize=11.5, leading=15, textColor=accent_text,
        spaceBefore=10, spaceAfter=7,
    )
    label = ParagraphStyle(
        "ResearchLabel", parent=small, fontSize=8, leading=10.5,
        textColor=colors.HexColor("#344054"),
        spaceAfter=2,
    )
    header_label = ParagraphStyle(
        "ResearchHeaderLabel", parent=label, fontSize=8.2, leading=10.5,
        textColor=colors.white,
    )
    def p(value, style=body, fallback="未填写"):
        return Paragraph(_pdf_value(value, fallback), style)

    def table(headers, rows, widths=None):
        if not rows:
            return p("暂无数据", small)
        table_data = [[p(header, header_label) for header in headers]]
        table_data.extend([
            [p(value, body, "—") for value in row]
            for row in rows
        ])
        if widths is None:
            widths = [document.width / len(headers)] * len(headers)
        result = Table(table_data, colWidths=widths, repeatRows=1, hAlign="LEFT")
        result.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), accent),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, -1), font_name),
            ("GRID", (0, 0), (-1, -1), 0.35, line),
            ("BACKGROUND", (0, 1), (-1, -1), colors.white),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7F9FA")]),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        return result

    def section_heading(text):
        return Paragraph(_pdf_value(text), section)

    def labelled_text(label_text, value, fill=None):
        content = [p(label_text, label), p(value, body)]
        if fill:
            box = Table([[content]], colWidths=[document.width])
            box.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), fill),
                ("BOX", (0, 0), (-1, -1), 0.4, line),
                ("LEFTPADDING", (0, 0), (-1, -1), 9),
                ("RIGHTPADDING", (0, 0), (-1, -1), 9),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            return [box]
        return content

    def image_preview(path, max_width=None, max_height=None):
        try:
            image_bytes = path.read_bytes()
            reader = ImageReader(io.BytesIO(image_bytes))
            width, height = reader.getSize()
            if not width or not height:
                return None
            max_width = max_width or min(document.width, 150 * mm)
            max_height = max_height or 78 * mm
            scale = min(max_width / width, max_height / height, 1)
            image = Image(io.BytesIO(image_bytes), width=width * scale, height=height * scale)
            image.hAlign = "LEFT"
            return image
        except Exception:
            return None

    def attachment_story(record_id, attachments):
        result = [section_heading("附件与原始数据")]
        preview_cells = []
        resolved_previews = _resolved_attachment_previews(
            attachments, attachment_lookup, attachment_path_resolver,
        )
        for _, source in resolved_previews:
            preview = image_preview(
                source,
                max_width=(72 if len(resolved_previews) > 1 else 96) * mm,
                max_height=56 * mm,
            )
            if preview:
                preview_cells.append([preview])
        if preview_cells:
            preview_table = Table(
                [preview_cells],
                colWidths=[document.width / len(preview_cells)] * len(preview_cells),
                hAlign="LEFT",
            )
            preview_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F7F9FA")),
                ("BOX", (0, 0), (-1, -1), 0.4, line),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, line),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]))
            result.extend([preview_table, Spacer(1, 5)])

        folder_name, folder_note = _attachment_folder_note(record_id, attachments)
        folder_box = Table(
            [[p("完整实验文件夹", label), p(f"{folder_name}  ·  {folder_note}", body)]],
            colWidths=[36 * mm, document.width - 36 * mm],
            hAlign="LEFT",
        )
        folder_box.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), soft),
            ("BOX", (0, 0), (-1, -1), 0.4, line),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ]))
        result.append(folder_box)
        return [KeepTogether(result)]

    story = [
        Spacer(1, 23 * mm),
        Paragraph(_pdf_value(profile["kicker"]), kicker),
        Paragraph(_pdf_value(experiment["title"], "实验记录"), cover),
        Paragraph(_pdf_value(
            f"{experiment.get('project_title') or '未归属项目'}  ·  {experiment.get('code') or '未设置编号'}"
        ), body),
        Spacer(1, 8 * mm),
        Table([[p("实验状态", label), p(experiment.get("status"), body),
                p("最后导出", label), p(payload["exported_at"].replace("T", " "), body)]],
              colWidths=[24 * mm, 46 * mm, 24 * mm, document.width - 94 * mm],
              style=TableStyle([
                  ("BACKGROUND", (0, 0), (-1, -1), soft),
                  ("BOX", (0, 0), (-1, -1), 0.4, line),
                  ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                  ("LEFTPADDING", (0, 0), (-1, -1), 8),
                  ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                  ("TOPPADDING", (0, 0), (-1, -1), 7),
                  ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
              ])),
        Spacer(1, 12 * mm),
        p("本报告由 R/LAB Research Assistant 生成。原始文件仍保存在实验资料目录中，报告仅展示核心缩略图和完整实验文件夹。", small),
        PageBreak(),
        section_heading("实验概览"),
        table(
            ("字段", "内容"),
            (("项目", experiment.get("project_title")), ("实验编号", experiment.get("code")),
             ("负责人", experiment.get("owner")), ("状态", experiment.get("status")),
             ("计划时间", f"{experiment.get('start_date') or '未设置'} - {experiment.get('end_date') or '未设置'}")),
            widths=[34 * mm, document.width - 34 * mm],
        ),
        Spacer(1, 5),
            *labelled_text("实验目的与背景", experiment.get("objective"), soft),
        section_heading("计划步骤"),
        table(
            ("序号", "步骤", "计划执行人", "计划日期", "说明"),
            [
                (step["position"], step["title"], step["operator"], step["planned_date"] or "未安排", step["description"])
                for step in payload["steps"]
            ],
            widths=[13 * mm, 37 * mm, 28 * mm, 28 * mm, document.width - 106 * mm],
        ),
        section_heading("计划参数"),
        table(
            ("序号", "参数", "数值", "单位", "说明"),
            [
                (item["position"], item["name"], item["value"], item["unit"], item["notes"])
                for item in payload["plan_parameters"]
            ],
            widths=[13 * mm, 38 * mm, 26 * mm, 20 * mm, document.width - 97 * mm],
        ),
        section_heading("关联样本"),
        table(
            ("编号", "类型", "用途", "使用量", "来源", "状态", "备注"),
            [
                (sample["sample_code"], sample["sample_type"], sample["role"], sample["amount_used"],
                 sample["source"], sample["status"], sample["notes"])
                for sample in payload["samples"]
            ],
            widths=[25 * mm, 24 * mm, 25 * mm, 22 * mm, 25 * mm, 22 * mm, document.width - 143 * mm],
        ),
    ]

    groups = execution_groups(payload)
    for execution_index, execution in enumerate(groups, start=1):
        if execution_index > 1:
            story.append(PageBreak())
        story.extend([
            section_heading(f"实验批次 {execution_index:02d} · {execution['batch_code']}"),
            table(
                ("字段", "内容"),
                (("重复类型", execution["repeat_kind"]), ("分组", execution["group_name"]),
                 ("实验人员", execution["operator"]), ("状态", execution["status"]),
                 ("实际时间", f"{execution['start_date'] or '未设置'} - {execution['end_date'] or '进行中'}"),
                 ("质量标记", "建议重复" if execution["requires_repeat"] else "无")),
                widths=[34 * mm, document.width - 34 * mm],
            ),
        ])
        story.extend(labelled_text("批次摘要", execution.get("summary"), soft))
        if execution.get("conclusion"):
            story.extend(labelled_text("批次结论", execution.get("conclusion"), colors.HexColor("#FFF4D6")))
        story.append(section_heading("批次步骤"))
        story.append(table(
            ("序号", "状态", "步骤", "执行人", "计划日期", "完成日期", "说明"),
            [
                (step["position"], "已完成" if step["is_done"] else "待完成", step["title"],
                 step["operator"], step["planned_date"] or "未安排", step["completed_date"] or "未完成",
                 step["description"])
                for step in execution["steps"]
            ],
            widths=[12 * mm, 19 * mm, 31 * mm, 24 * mm, 25 * mm, 25 * mm, document.width - 136 * mm],
        ))
        if execution["actual_parameters"]:
            story.extend([section_heading("批次参数"), table(
                ("参数", "数值", "单位", "说明"),
                [(item["name"], item["value"], item["unit"], item["notes"]) for item in execution["actual_parameters"]],
                widths=[42 * mm, 30 * mm, 23 * mm, document.width - 95 * mm],
            )])
        story.append(section_heading("过程记录"))
        if not execution["records"]:
            story.append(p("本批次尚无过程记录。", small))
        for record_index, record in enumerate(execution["records"], start=1):
            record_story = [
                Paragraph(_pdf_value(
                    f"{execution_index:02d}.{record_index:02d}  {record['record_date'] or '未填写日期'}  ·  {record['result']}"
                ), title),
                table(
                    ("字段", "内容"),
                    (("实验人员", record["operator"]), ("记录状态", record["lifecycle_status"])),
                    widths=[34 * mm, document.width - 34 * mm],
                ),
            ]
            record_story.extend(labelled_text("实验条件", record.get("conditions"), soft))
            record_story.extend(labelled_text("操作与观察", record.get("content")))
            record_story.extend(labelled_text("结论与备注", record.get("remark"), colors.HexColor("#FFF4D6")))
            story.append(KeepTogether(record_story))
            if record["parameters"]:
                story.append(table(
                    ("参数", "数值", "单位", "说明"),
                    [(item["name"], item["value"], item["unit"], item["notes"]) for item in record["parameters"]],
                    widths=[42 * mm, 30 * mm, 23 * mm, document.width - 95 * mm],
                ))
            story.extend(attachment_story(record["id"], record["attachments"]))

    def on_page(canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(line)
        canvas.setLineWidth(0.5)
        canvas.line(doc.leftMargin, 12 * mm, A4[0] - doc.rightMargin, 12 * mm)
        canvas.setFont(font_name, 7.5)
        canvas.setFillColor(muted)
        canvas.drawString(doc.leftMargin, 7 * mm, f"R/LAB  ·  {profile['label']}")
        canvas.drawRightString(A4[0] - doc.rightMargin, 7 * mm, f"第 {doc.page} 页")
        canvas.restoreState()

    document.build(story, onFirstPage=on_page, onLaterPages=on_page)
    return output.getvalue()


def _record_report_payload(record):
    experiment = record.experiment
    batch = record.batch
    return {
        "experiment_title": experiment.title,
        "experiment_code": experiment.code,
        "project_title": experiment.project.title if experiment.project else "",
        "objective": experiment.objective,
        "record_id": record.id,
        "record_date": _date_value(record.record_date),
        "operator": record.operator,
        "batch_code": batch.batch_code if batch else "历史记录",
        "batch_status": batch.status if batch else "待归档",
        "conditions": record.conditions,
        "content": record.content,
        "result": record.result,
        "lifecycle_status": record.lifecycle_status,
        "remark": record.remark,
        "parameters": [_parameter_payload(parameter) for parameter in record.parameters],
        "attachments": [
            _attachment_payload(attachment)
            for attachment in _active(record.attachments)
        ],
    }


def build_record_docx_export(record, attachment_path_resolver=None, template_key="research"):
    """Render one experiment record in the same table layout as the report reader."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    profile = resolve_report_template(template_key)
    payload = _record_report_payload(record)
    attachment_lookup = {
        attachment.id: attachment
        for attachment in _active(record.attachments)
    }
    document = Document()
    _set_docx_font(document, profile)
    document.core_properties.title = f"{payload['experiment_title']} - 实验记录"
    document.core_properties.subject = "单条实验记录报告"
    document.core_properties.author = "R/LAB Research Assistant"

    title = document.add_heading(payload["experiment_title"], 0)
    title.paragraph_format.space_after = Pt(3)
    subtitle = document.add_paragraph("EXPERIMENT RECORD  /  实验记录报告")
    subtitle.runs[0].font.size = Pt(9)
    subtitle.runs[0].font.bold = True
    subtitle.runs[0].font.color.rgb = RGBColor.from_string(profile["text_accent"])
    subtitle.paragraph_format.space_after = Pt(13)

    _add_docx_table(profile, document, ("字段", "内容"), (
        ("实验名称", payload["experiment_title"]),
        ("项目", payload["project_title"]),
        ("实验编号", payload["experiment_code"] or "未设置"),
        ("记录编号", f"R-{payload['record_id']}"),
        ("实验日期", payload["record_date"]),
        ("所属批次", payload["batch_code"]),
        ("实验人员", payload["operator"]),
        ("结果判定", payload["result"]),
        ("记录状态", payload["lifecycle_status"]),
    ))
    _add_docx_callout(profile, document, "实验目的", payload["objective"], "EDF3FF")
    _add_docx_callout(profile, document, "实验条件 / 仪器设备 / 原理", payload["conditions"], "F3F6F7")
    _add_docx_callout(profile, document, "实验过程与观察记录", payload["content"], "FFFFFF")
    _add_docx_callout(profile, document, "实验结果 / 结论 / 备注", payload["remark"], "FFF4D6")

    document.add_heading("结构化参数", level=1)
    _add_docx_table(profile, document, ("参数", "数值", "单位", "说明"), [
        (parameter["name"], parameter["value"], parameter["unit"], parameter["notes"])
        for parameter in payload["parameters"]
    ])

    document.add_heading("原始数据与附件", level=1)
    _add_docx_attachment_section(
        profile, document, payload["record_id"], payload["attachments"], attachment_lookup,
        attachment_path_resolver,
    )
    review = document.add_paragraph(
        "人工核验提示：自动导入或整理的内容只作为记录草稿，请结合原始文件和实验室 SOP 复核。"
    )
    review.runs[0].italic = True
    review.runs[0].font.size = Pt(8.5)
    review.runs[0].font.color.rgb = RGBColor.from_string("52616B")
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.LEFT

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def build_record_pdf_export(record, template_key="research", attachment_path_resolver=None):
    """Render one experiment record as a paginated, table-based PDF report."""
    try:
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.lib.utils import ImageReader
        from reportlab.platypus import KeepTogether, Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError as exc:
        raise RuntimeError("PDF 导出需要安装 reportlab 依赖。") from exc

    profile = REPORT_TEMPLATES.get(template_key, REPORT_TEMPLATES["research"])
    payload = _record_report_payload(record)
    attachment_lookup = {
        attachment.id: attachment
        for attachment in _active(record.attachments)
    }
    font_name = _pdf_font_name()
    accent = colors.HexColor(f"#{profile['accent']}")
    accent_text = colors.HexColor(f"#{profile['text_accent']}")
    soft = colors.HexColor(f"#{profile['soft']}")
    ink = colors.HexColor("#17212B")
    muted = colors.HexColor("#475467")
    line = colors.HexColor("#CBD5DC")
    output = io.BytesIO()
    document = SimpleDocTemplate(
        output, pagesize=A4, rightMargin=16 * mm, leftMargin=16 * mm,
        topMargin=17 * mm, bottomMargin=17 * mm,
        title=f"{payload['experiment_title']} - 实验记录",
        author="R/LAB Research Assistant",
    )
    body = ParagraphStyle("RecordBody", fontName=font_name, fontSize=9.4, leading=14.2, textColor=ink, spaceAfter=4)
    small = ParagraphStyle("RecordSmall", parent=body, fontSize=8, leading=11, textColor=muted)
    title = ParagraphStyle("RecordTitle", parent=body, fontSize=20, leading=25, textColor=ink, spaceAfter=4)
    kicker = ParagraphStyle("RecordKicker", parent=small, fontSize=8.2, leading=10.5, textColor=accent_text, spaceAfter=7)
    section = ParagraphStyle("RecordSection", parent=body, fontSize=11.5, leading=15, textColor=accent_text, spaceBefore=9, spaceAfter=6)
    label = ParagraphStyle(
        "RecordLabel", parent=small, fontSize=8, leading=10.5,
        textColor=colors.HexColor("#344054"),
    )
    header_label = ParagraphStyle(
        "RecordHeaderLabel", parent=label, fontSize=8.2, leading=10.5,
        textColor=colors.white,
    )

    def p(value, style=body, fallback="未填写"):
        return Paragraph(_pdf_value(value, fallback), style)

    def table(headers, rows, widths=None):
        data = [[p(header, header_label) for header in headers]]
        data.extend([[p(value, body, "-") for value in row] for row in rows])
        result = Table(data, colWidths=widths or [document.width / len(headers)] * len(headers), repeatRows=1, hAlign="LEFT")
        result.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), accent),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, -1), font_name),
            ("GRID", (0, 0), (-1, -1), 0.35, line),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F7F9FA")]),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 6),
            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        return result

    def section_heading(value):
        return Paragraph(_pdf_value(value), section)

    def field_box(label_text, value, fill=colors.white):
        box = Table([[p(label_text, label), p(value, body)]], colWidths=[42 * mm, document.width - 42 * mm])
        box.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), fill),
            ("BOX", (0, 0), (-1, -1), 0.35, line),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 7),
            ("RIGHTPADDING", (0, 0), (-1, -1), 7),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]))
        return box

    def image_preview(path, max_width, max_height):
        try:
            image_bytes = path.read_bytes()
            reader = ImageReader(io.BytesIO(image_bytes))
            width, height = reader.getSize()
            if not width or not height:
                return None
            scale = min(max_width / width, max_height / height, 1)
            image = Image(io.BytesIO(image_bytes), width=width * scale, height=height * scale)
            image.hAlign = "LEFT"
            return image
        except Exception:
            return None

    def attachment_story():
        result = [section_heading("原始数据与附件")]
        resolved_previews = _resolved_attachment_previews(
            payload["attachments"], attachment_lookup, attachment_path_resolver,
        )
        preview_cells = []
        for _, source in resolved_previews:
            preview = image_preview(
                source,
                (72 if len(resolved_previews) > 1 else 96) * mm,
                56 * mm,
            )
            if preview:
                preview_cells.append([preview])
        if preview_cells:
            preview_table = Table(
                [preview_cells],
                colWidths=[document.width / len(preview_cells)] * len(preview_cells),
                hAlign="LEFT",
            )
            preview_table.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#F7F9FA")),
                ("BOX", (0, 0), (-1, -1), 0.4, line),
                ("INNERGRID", (0, 0), (-1, -1), 0.35, line),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 7),
                ("RIGHTPADDING", (0, 0), (-1, -1), 7),
                ("TOPPADDING", (0, 0), (-1, -1), 7),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
            ]))
            result.extend([preview_table, Spacer(1, 5)])

        folder_name, folder_note = _attachment_folder_note(
            payload["record_id"], payload["attachments"],
        )
        folder_box = Table(
            [[p("完整实验文件夹", label), p(f"{folder_name}  ·  {folder_note}", body)]],
            colWidths=[42 * mm, document.width - 42 * mm],
            hAlign="LEFT",
        )
        folder_box.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, -1), soft),
            ("BOX", (0, 0), (-1, -1), 0.4, line),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("LEFTPADDING", (0, 0), (-1, -1), 8),
            ("RIGHTPADDING", (0, 0), (-1, -1), 8),
            ("TOPPADDING", (0, 0), (-1, -1), 7),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ]))
        result.append(folder_box)
        return [KeepTogether(result)]

    story = [
        Paragraph(_pdf_value(profile["kicker"]), kicker),
        Paragraph(_pdf_value(payload["experiment_title"], "实验记录"), title),
        Paragraph(_pdf_value(f"{payload['project_title'] or '未归属项目'}  ·  {payload['batch_code']}"), body),
        Spacer(1, 5 * mm),
        table(("字段", "内容"), [
            ("实验编号", payload["experiment_code"] or "未设置"),
            ("记录编号", f"R-{payload['record_id']}"),
            ("实验日期", payload["record_date"]),
            ("实验人员", payload["operator"]),
            ("所属批次", payload["batch_code"]),
            ("结果判定", payload["result"]),
            ("记录状态", payload["lifecycle_status"]),
        ], [42 * mm, document.width - 42 * mm]),
        section_heading("实验目的"), field_box("研究目的", payload["objective"], soft),
        section_heading("实验条件与过程"), field_box("条件 / 仪器 / 原理", payload["conditions"], colors.HexColor("#F3F6F7")),
        field_box("过程与观察记录", payload["content"]),
        section_heading("实验结果"), field_box("结果 / 结论 / 备注", payload["remark"], colors.HexColor("#FFF4D6")),
        section_heading("结构化参数"),
        table(("参数", "数值", "单位", "说明"), [
            (parameter["name"], parameter["value"], parameter["unit"], parameter["notes"])
            for parameter in payload["parameters"]
        ], [42 * mm, 30 * mm, 23 * mm, document.width - 95 * mm]),
        *attachment_story(),
        Spacer(1, 4 * mm),
        p("人工核验提示：自动导入或整理的内容只作为记录草稿，请结合原始文件和实验室 SOP 复核。", small),
    ]

    def on_page(canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(line)
        canvas.setLineWidth(0.5)
        canvas.line(doc.leftMargin, 12 * mm, A4[0] - doc.rightMargin, 12 * mm)
        canvas.setFont(font_name, 7.5)
        canvas.setFillColor(muted)
        canvas.drawString(doc.leftMargin, 7 * mm, "R/LAB  ·  实验记录报告")
        canvas.drawRightString(A4[0] - doc.rightMargin, 7 * mm, f"第 {doc.page} 页")
        canvas.restoreState()

    document.build(story, onFirstPage=on_page, onLaterPages=on_page)
    return output.getvalue()


def _add_xlsx_sheet(workbook, title, headers, rows):
    from openpyxl.styles import Alignment, Font, PatternFill
    from openpyxl.utils import get_column_letter

    sheet = workbook.create_sheet(title)
    sheet.append(list(headers))
    for row in rows:
        sheet.append([_text(value) for value in row])
    header_fill = PatternFill("solid", fgColor="2166F3")
    for cell in sheet[1]:
        cell.fill = header_fill
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(vertical="center")
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    for column_index, header in enumerate(headers, start=1):
        values = [header, *[row[column_index - 1] for row in rows if len(row) >= column_index]]
        width = min(max(max((len(_text(value)) for value in values), default=8) + 2, 10), 48)
        sheet.column_dimensions[get_column_letter(column_index)].width = width
    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    return sheet


def build_xlsx_export(item):
    from openpyxl import Workbook

    payload = experiment_payload(item)
    experiment = payload["experiment"]
    groups = execution_groups(payload)
    workbook = Workbook()
    workbook.remove(workbook.active)
    _add_xlsx_sheet(workbook, "实验信息", ("字段", "内容"), [
        ("实验名称", experiment["title"]), ("实验编号", experiment["code"]),
        ("状态", experiment["status"]), ("负责人", experiment["owner"]),
        ("计划开始", experiment["start_date"]), ("计划结束", experiment["end_date"]),
        ("实验目的", experiment["objective"]), ("导出时间", payload["exported_at"]),
    ])
    _add_xlsx_sheet(workbook, "关联样本", ("样本编号", "类型", "来源", "位置", "库存量", "状态", "用途", "使用量", "备注"), [
        (sample["sample_code"], sample["sample_type"], sample["source"], sample["location"],
         sample["quantity"], sample["status"], sample["role"], sample["amount_used"], sample["notes"])
        for sample in payload["samples"]
    ])
    _add_xlsx_sheet(workbook, "计划参数", ("序号", "参数", "数值", "单位", "说明"), [
        (parameter["position"], parameter["name"], parameter["value"], parameter["unit"], parameter["notes"])
        for parameter in payload["plan_parameters"]
    ])
    _add_xlsx_sheet(workbook, "实验步骤", ("序号", "步骤", "计划执行人", "计划日期", "说明"), [
        (step["position"], step["title"], step["operator"], step["planned_date"], step["description"])
        for step in payload["steps"]
    ])
    _add_xlsx_sheet(workbook, "实验批次", (
        "批次 ID", "批次编号", "重复类型", "重复序号", "实验分组", "实验人员", "状态",
        "实际开始", "实际结束", "批次摘要", "批次结论", "建议重复", "记录数",
    ), [
        (execution["id"], execution["batch_code"], execution["repeat_kind"], execution["repeat_number"],
         execution["group_name"], execution["operator"], execution["status"], execution["start_date"],
         execution["end_date"], execution["summary"], execution["conclusion"],
         execution["requires_repeat"], len(execution["records"]))
        for execution in groups
    ])
    _add_xlsx_sheet(workbook, "批次步骤", (
        "批次 ID", "批次编号", "步骤 ID", "来源计划步骤 ID", "序号", "状态", "步骤",
        "实验人员", "计划日期", "完成日期", "说明",
    ), [
        (
            execution["id"], execution["batch_code"], step["id"], step["source_step_id"],
            step["position"], "已完成" if step["is_done"] else "待完成", step["title"],
            step["operator"], step["planned_date"], step["completed_date"], step["description"],
        )
        for execution in groups for step in execution["steps"]
    ])
    _add_xlsx_sheet(workbook, "过程记录", (
        "批次 ID", "批次编号", "过程记录 ID", "日期", "实验人员", "结果", "实验条件", "实验过程", "结论与备注",
    ), [
        (record["batch_id"], record["batch_code"], record["id"], record["record_date"],
         record["operator"], record["result"], record["conditions"], record["content"], record["remark"])
        for record in payload["records"]
    ])
    _add_xlsx_sheet(workbook, "过程记录参数", (
        "批次编号", "过程记录 ID", "日期", "序号", "参数", "数值", "单位", "说明",
    ), [
        (record["batch_code"], record["id"], record["record_date"], parameter["position"], parameter["name"],
         parameter["value"], parameter["unit"], parameter["notes"])
        for record in payload["records"] for parameter in record["parameters"]
    ])
    _add_xlsx_sheet(workbook, "附件清单", (
        "批次编号", "过程记录 ID", "日期", "分类", "文件", "版本", "大小（字节）", "类型", "SHA-256", "标签", "说明",
    ), [
        (record["batch_code"], record["id"], record["record_date"], attachment["category"], attachment["relative_path"],
         attachment["version_number"], attachment["size_bytes"], attachment["mime_type"],
         attachment["sha256"], attachment["tags"], attachment["description"])
        for record in payload["records"] for attachment in record["attachments"]
    ])
    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def _archive_component(value, fallback):
    cleaned = re.sub(r'[\\/:*?"<>|\x00-\x1f]', "_", _text(value).strip()).strip(". ")
    return cleaned[:120] or fallback


def _archive_relative_parts(value):
    parts = []
    for index, raw_part in enumerate(_text(value).replace("\\", "/").split("/"), start=1):
        if raw_part in {"", "."}:
            continue
        if raw_part == "..":
            raw_part = "parent"
        parts.append(_archive_component(raw_part, f"item-{index}"))
    return parts or ["file"]


def build_archive_export(item, attachment_path_resolver):
    """Build a ZIP from the same filtered payload used by every other export."""
    payload = experiment_payload(item)
    record_by_id = {record.id: record for record in _active(item.records)}
    attachment_by_id = {
        attachment.id: attachment
        for record in record_by_id.values()
        for attachment in _active(record.attachments)
    }
    archive = tempfile.SpooledTemporaryFile(max_size=8 * 1024 * 1024, mode="w+b")
    manifest_output = io.StringIO()
    manifest_writer = csv.writer(manifest_output)
    manifest_writer.writerow([
        "execution_id", "execution_code", "record_date", "record_id", "category", "relative_path",
        "version", "size_bytes", "mime_type", "sha256", "tags", "description", "archive_path",
    ])
    with zipfile.ZipFile(
        archive, "w", compression=zipfile.ZIP_DEFLATED, allowZip64=True
    ) as bundle:
        bundle.writestr("report.md", _markdown_from_payload(payload).encode("utf-8"))
        bundle.writestr(
            "experiment.json",
            json.dumps(payload, ensure_ascii=False, indent=2).encode("utf-8"),
        )
        for execution in execution_groups(payload):
            execution_code = execution["batch_code"]
            execution_folder = _archive_component(execution_code, "execution")
            for record in execution["records"]:
                for attachment_data in record["attachments"]:
                    attachment = attachment_by_id.get(attachment_data["id"])
                    path_parts = _archive_relative_parts(attachment_data["relative_path"])
                    if attachment_data["version_number"] > 1:
                        path_parts[-1] = f"v{attachment_data['version_number']}-{path_parts[-1]}"
                    archive_path = "/".join([
                        "files",
                        execution_folder,
                        record["record_date"] or "unknown-date",
                        f"record-{record['id']}",
                        _archive_component(attachment_data["category"], "uncategorized"),
                        *path_parts,
                    ])
                    source_path = attachment_path_resolver(attachment) if attachment else None
                    exists = bool(source_path and source_path.is_file())
                    if exists:
                        bundle.write(source_path, archive_path)
                    elif attachment_data["storage_mode"] == "external":
                        archive_path = "外部链接（未打包）"
                    else:
                        archive_path = "文件缺失"
                    manifest_writer.writerow([
                        record["batch_id"], execution_code, record["record_date"], record["id"],
                        attachment_data["category"], attachment_data["relative_path"],
                        attachment_data["version_number"], attachment_data["size_bytes"],
                        attachment_data["mime_type"], attachment_data["sha256"], attachment_data["tags"],
                        attachment_data["description"], archive_path,
                    ])
        bundle.writestr("file-manifest.csv", ("\ufeff" + manifest_output.getvalue()).encode("utf-8"))
    archive.seek(0)
    return archive
