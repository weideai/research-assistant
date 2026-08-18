import json
from pathlib import Path

import pytest
from sqlalchemy import inspect

from app import db
from app.models import (
    AIChangeSet, AIConversation, AIMessage, ApiPreset, CalendarEvent, LabRecordRevision,
    LibraryItem, SearchDocument, Task, User, WeeklyReport, WeeklyReportUpdate,
)
from app.services.desktop_modules import AI_TARGET_FIELDS, build_desktop_ai_system_prompt
from app.services.desktop_workspace import ConflictError, DesktopApplicationService


def desktop_service(app):
    with app.app_context():
        user = User(
            name="本地研究者", email="complete-desktop@example.invalid",
            password_hash="local", role="system_admin",
        )
        db.session.add(user)
        db.session.commit()
    return DesktopApplicationService(app)


def test_clean_desktop_database_bootstraps_local_user_and_workspace(app):
    service = DesktopApplicationService(app)
    dashboard = service.dashboard()
    assert dashboard["workspace"]["name"] == "R/LAB 工作区"
    with app.app_context():
        assert User.query.filter_by(email="local@research-assistant.invalid").count() == 1


def test_complete_resource_planning_and_search_workflow(app):
    service = desktop_service(app)
    project = service.create_project({"title": "蛋白表达优化", "code": "PRO-01"})
    record = service.create_record({"project_id": project["id"], "title": "IPTG 梯度实验", "experiment_date": "2026-08-11"})

    literature = service.save_literature({
        "title": "Protein expression optimization", "authors": "Li; Wang",
        "year": 2025, "doi": "https://doi.org/10.1000/Example", "read_status": "reading",
    })
    linked = service.link_literature({"literature_id": literature["id"], "project_id": project["id"]})
    assert linked["project_ids"] == [project["id"]]
    assert linked["doi"] == "https://doi.org/10.1000/Example"

    note = service.save_note({
        "title": "诱导条件观察要点", "kind": "experiment_guide",
        "project_id": project["id"], "body": "记录菌液浊度和表达条带。",
    })
    saved_note = service.save_note({"id": note["id"], "title": note["title"], "kind": note["kind"], "body": "补充 OD600。"}, note["row_version"])
    assert saved_note["row_version"] == 2
    with pytest.raises(ConflictError):
        service.save_note({"id": note["id"], "title": "过期修改", "kind": "general"}, note["row_version"])

    task = service.save_task({
        "title": "整理 IPTG 梯度数据", "project_id": project["id"],
        "lab_record_id": record["id"], "deadline": "2026-08-12", "priority": "high",
    })
    completed = service.save_task({**task, "status": "done"}, task["row_version"])
    assert completed["completed_at"] is not None
    assert service.list_tasks({"scope": "completed"})[0]["id"] == task["id"]

    current = service.weekly_current({"week": "2026-08-11"})
    report = service.save_weekly({
        "period_start": current["period_start"], "body": "完成诱导条件预实验。",
        "issues_and_feedback": "需要补充重复。", "next_week_plan": "完成蛋白定量。",
        "entries": current["entries"],
    })
    assert report["row_version"] == 1
    assert service.list_weekly()[0]["id"] == report["id"]

    results = service.search({"query": "IPTG"})
    assert any(item["entity_type"] == "record" for item in results["items"])

    service.trash_move({"entity_type": "note", "id": note["id"]})
    assert service.trash_list()[0]["entity_type"] == "note"
    service.trash_restore({"entity_type": "note", "id": note["id"]})
    assert service.get_note({"id": note["id"]})["title"] == note["title"]


def test_ai_change_set_only_applies_accepted_fields_and_creates_revision(app):
    service = desktop_service(app)
    project = service.create_project({"title": "AI 审计项目"})
    record = service.create_record({"project_id": project["id"], "title": "AI 审计记录"})

    with app.app_context():
        user_id = User.query.filter_by(email="complete-desktop@example.invalid").one().id
        conversation = AIConversation(user_id=user_id, title="字段建议")
        db.session.add(conversation)
        db.session.flush()
        message = AIMessage(conversation_id=conversation.id, role="assistant", content="建议")
        db.session.add(message)
        db.session.flush()
        change = AIChangeSet(
            message_id=message.id, target_type="lab_record", target_id=record["id"],
            base_row_version=record["row_version"],
            proposal_json=json.dumps({"objective": "新的实验目的", "conclusion": "不应写入"}, ensure_ascii=False),
            before_json="{}", source_snapshot_json="{}",
        )
        db.session.add(change)
        db.session.commit()
        change_id = change.id

    applied = service.ai_apply({"id": change_id, "accepted_fields": ["objective"]})
    assert applied["status"] == "partially_applied"
    updated = service.get_record({"id": record["id"]})
    assert updated["objective"] == "新的实验目的"
    assert updated["conclusion"] == ""
    with app.app_context():
        assert LabRecordRevision.query.filter_by(source_ai_change_set_id=change_id).count() == 1


def test_ai_writes_tasks_calendar_and_weekly_annotations_without_risk_confirmation(app):
    service = desktop_service(app)
    project = service.create_project({"title": "AI 调度项目"})
    task = service.save_task({"title": "原始任务", "project_id": project["id"]})
    event = service.create_calendar_event({
        "title": "原始日历事件", "starts_at": "2026-08-12T09:00:00",
        "project_id": project["id"],
    })
    report = service.save_weekly({
        "title": "AI 周报", "period_start": "2026-08-10", "body": "本周正文", "entries": [],
    })

    proposals = [
        ("task", task["id"], task["row_version"], {"title": "整理实验数据", "priority": "high", "deadline": "2026-08-15"}),
        ("calendar_event", event["id"], event["row_version"], {"title": "课题讨论", "starts_at": "2026-08-13T10:30:00", "notes": "讨论下一轮实验"}),
        ("weekly_report", report["id"], report["row_version"], {"annotation": "补充阴性对照并说明重复次数。"}),
    ]
    change_ids = []
    with app.app_context():
        user_id = User.query.filter_by(email="complete-desktop@example.invalid").one().id
        conversation = AIConversation(user_id=user_id, title="多目标写入")
        db.session.add(conversation)
        db.session.flush()
        for target_type, target_id, row_version, proposal in proposals:
            message = AIMessage(conversation_id=conversation.id, role="assistant", content="建议")
            db.session.add(message)
            db.session.flush()
            change = AIChangeSet(
                message_id=message.id, target_type=target_type, target_id=target_id,
                base_row_version=row_version, proposal_json=json.dumps(proposal, ensure_ascii=False),
                before_json="{}", source_snapshot_json="{}",
            )
            db.session.add(change)
            db.session.flush()
            change_ids.append(change.id)
        db.session.commit()

    task_change, event_change, report_change = change_ids
    service.ai_apply({"id": task_change, "accepted_fields": list(proposals[0][3])})
    service.ai_apply({"id": event_change, "accepted_fields": list(proposals[1][3])})
    service.ai_apply({"id": report_change, "accepted_fields": ["annotation"]})

    saved_task = next(item for item in service.list_tasks() if item["id"] == task["id"])
    assert saved_task["title"] == "整理实验数据"
    assert saved_task["priority"] == "high"
    assert saved_task["deadline"] == "2026-08-15"
    with app.app_context():
        saved_event = db.session.get(CalendarEvent, event["id"])
        assert saved_event.title == "课题讨论"
        assert saved_event.starts_at.isoformat() == "2026-08-13T10:30:00"
        assert db.session.get(Task, task["id"]).row_version == task["row_version"] + 1
        assert WeeklyReportUpdate.query.filter_by(report_id=report["id"], kind="AI批注").count() == 1

    service.ai_revert({"id": report_change})
    with app.app_context():
        assert WeeklyReportUpdate.query.filter_by(report_id=report["id"], kind="AI批注").count() == 0


def test_desktop_ai_prompts_match_current_targets_and_experiment_report_fields():
    for target_type in AI_TARGET_FIELDS:
        schema = {field: {"type": "text"} for field in AI_TARGET_FIELDS[target_type]}
        prompt = build_desktop_ai_system_prompt(
            target_type, schema, today=__import__("datetime").date(2026, 8, 11), timezone="Asia/Shanghai",
        )
        assert f"目标类型：{target_type}" in prompt
        assert "直接写入" in prompt
        assert "experiment_plan" not in prompt
        assert "batch" not in prompt
        assert "实验计划" not in prompt
        assert "实验批次" not in prompt

    record_prompt = build_desktop_ai_system_prompt(
        "lab_record", {field: {"type": "text"} for field in AI_TARGET_FIELDS["lab_record"]},
        today=__import__("datetime").date(2026, 8, 11), timezone="Asia/Shanghai",
    )
    for section in (
        "研究背景", "实验目的", "实验假设", "实验设计", "材料与条件", "预期结果",
        "实际过程", "实际结果", "分析", "结论", "下一步",
    ):
        assert section in record_prompt
    assert "不能把预期结果改写成实际结果" in record_prompt

    weekly_prompt = build_desktop_ai_system_prompt(
        "weekly_report", {field: {"type": "text"} for field in AI_TARGET_FIELDS["weekly_report"]},
        today=__import__("datetime").date(2026, 8, 11), timezone="Asia/Shanghai",
    )
    assert "annotation 表示新增一条周报批注" in weekly_prompt


def test_ai_proposal_uses_target_prompt_and_disables_risk_review(app, monkeypatch):
    service = desktop_service(app)
    project = service.create_project({"title": "AI 提示词项目"})
    record = service.create_record({"project_id": project["id"], "title": "报告生成记录"})
    captured = {}

    with app.app_context():
        user = User.query.filter_by(email="complete-desktop@example.invalid").one()
        preset = ApiPreset(
            user_id=user.id, name="测试模型", api_url="https://api.example.test/v1",
            text_model="test-model", is_enabled=True, is_default=True,
        )
        preset.set_api_key("test-key")
        db.session.add(preset)
        db.session.commit()

    def fake_chat(messages, system_prompt, config):
        captured.update(messages=messages, system_prompt=system_prompt, config=config)
        return {"reply": "已按当前实验记录生成。", "proposal": {"objective": "验证处理效果"}}

    monkeypatch.setattr("app.ai_service.chat_with_assistant", fake_chat)
    preview = service.ai_preview({"target_type": "lab_record", "target_id": record["id"]})
    result = service.ai_propose({
        "target_type": "lab_record", "target_id": record["id"],
        "prompt": "生成实验报告", "source_snapshot": preview["source_snapshot"],
    })

    assert result["proposal"] == {"objective": "验证处理效果"}
    assert "当前对象是一条完整实验记录" in captured["system_prompt"]
    assert "字段 Schema" in captured["system_prompt"]
    with app.app_context():
        message = AIMessage.query.order_by(AIMessage.id.desc()).first()
        assert message.requires_human_review is False


def test_ai_prompt_history_is_searchable_and_paginated(app):
    service = desktop_service(app)
    with app.app_context():
        user_id = User.query.filter_by(email="complete-desktop@example.invalid").one().id
        for index in range(7):
            conversation = AIConversation(
                user_id=user_id, title=f"历史 {index}", page_type="lab_record", page_id=42,
            )
            db.session.add(conversation)
            db.session.flush()
            message = AIMessage(
                conversation_id=conversation.id,
                role="assistant",
                content=f"回复 {index}",
                prompt_snapshot=f"生成实验报告 {index}",
                model_name="test-model",
                requires_human_review=False,
            )
            db.session.add(message)
            db.session.flush()
            db.session.add(AIChangeSet(
                message_id=message.id,
                target_type="lab_record",
                target_id=42,
                base_row_version=1,
                proposal_json="{}",
                before_json="{}",
                source_snapshot_json="{}",
                prompt_snapshot=message.prompt_snapshot,
                model_name=message.model_name,
            ))
        db.session.commit()

    first = service.ai_history({"page": 1, "per_page": 5, "query": "实验报告"})
    second = service.ai_history({"page": 2, "per_page": 5, "query": "实验报告"})

    assert first["pagination"] == {"page": 1, "pages": 2, "per_page": 5, "total": 7}
    assert len(first["items"]) == 5
    assert len(second["items"]) == 2
    assert all(item["target_type"] == "lab_record" for item in first["items"])
    assert all("实验报告" in item["prompt"] for item in first["items"])


def test_project_ai_context_reads_records_without_library_files(app, tmp_path):
    service = desktop_service(app)
    project = service.create_project({"title": "项目上下文", "objective": "理解完整研究历史"})
    record = service.create_record({"project_id": project["id"], "title": "无附件实验记录"})
    service.update_record({
        **record,
        "objective": "读取实验文本",
        "actual_result": "结果文本可以进入上下文。",
    }, record["row_version"])
    private_file = tmp_path / "must-not-enter-ai-context.txt"
    private_file.write_text("FILE-CONTENT-MUST-NOT-LEAK", encoding="utf-8")
    service.import_library_item({
        "path": str(private_file), "display_name": private_file.name,
        "storage_mode": "external", "kind": "raw_data", "project_id": project["id"],
    })

    preview = service.ai_preview({"target_type": "project", "target_id": project["id"]})
    snapshot = preview["source_snapshot"]

    assert snapshot["target_type"] == "project"
    assert snapshot["project_records"][0]["title"] == "无附件实验记录"
    assert snapshot["project_records"][0]["actual_result"] == "结果文本可以进入上下文。"
    serialized = json.dumps(snapshot, ensure_ascii=False)
    assert private_file.name not in serialized
    assert "FILE-CONTENT-MUST-NOT-LEAK" not in serialized
    assert "files" not in snapshot

    selected = service.ai_preview({
        "target_type": "project", "target_id": project["id"],
        "source_ids": [record["id"]], "records_explicit": True,
    })["source_snapshot"]
    empty = service.ai_preview({
        "target_type": "project", "target_id": project["id"],
        "source_ids": [], "records_explicit": True,
    })["source_snapshot"]
    assert [item["id"] for item in selected["project_records"]] == [record["id"]]
    assert empty["project_records"] == []


def test_project_ai_conversations_support_search_pagination_and_bulk_edit(app):
    service = desktop_service(app)
    project = service.create_project({"title": "聊天项目"})
    created = [
        service.ai_conversation_create({
            "project_id": project["id"], "record_ids": [], "title": f"阶段复盘 {index}",
        })
        for index in range(7)
    ]

    first = service.ai_conversations({
        "project_id": project["id"], "query": "阶段复盘", "page": 1, "per_page": 5,
    })
    second = service.ai_conversations({
        "project_id": project["id"], "query": "阶段复盘", "page": 2, "per_page": 5,
    })
    renamed = service.ai_conversations_bulk({
        "ids": [created[0]["id"], created[1]["id"]], "action": "rename", "title": "批量整理",
    })
    deleted = service.ai_conversations_bulk({
        "ids": [created[2]["id"], created[3]["id"]], "action": "delete",
    })

    assert first["pagination"]["total"] == 7
    assert len(first["items"]) == 5
    assert len(second["items"]) == 2
    assert renamed == {"updated": 2, "skipped": 0}
    assert deleted == {"updated": 2, "skipped": 0}
    remaining = service.ai_conversations({"project_id": project["id"], "page": 1, "per_page": 10})
    assert remaining["pagination"]["total"] == 5
    assert sum(item["title"].startswith("批量整理") for item in remaining["items"]) == 2


def test_weekly_upload_management_accepts_manual_annotations(app, tmp_path):
    service = desktop_service(app)
    project = service.create_project({"title": "周报批注项目"})
    report_file = tmp_path / "week-33.pdf"
    report_file.write_bytes(b"weekly report")
    report = service.import_weekly_file({
        "path": str(report_file), "title": "第 33 周总结", "report_date": "2026-08-12",
        "project_id": project["id"], "storage_mode": "managed", "status": "submitted",
    })

    annotated = service.add_weekly_annotation({
        "id": report["id"], "kind": "指导", "status": "待处理",
        "content": "补充下周实验的重复次数。",
    }, report["row_version"])

    assert annotated["annotation_count"] == 1
    assert annotated["updates"][0]["kind"] == "指导"
    assert annotated["updates"][0]["content"] == "补充下周实验的重复次数。"
    assert annotated["row_version"] == report["row_version"] + 1


def test_ai_can_replace_and_revert_structured_lab_record_steps(app):
    service = desktop_service(app)
    project = service.create_project({"title": "步骤写入项目"})
    record = service.create_record({"project_id": project["id"], "title": "步骤写入记录"})
    original = service.update_record({
        **record,
        "steps": [{
            "title": "原步骤", "instruction": "原操作", "planned_duration_minutes": 15,
            "checkpoint": "原检查点", "risk": "原注意事项", "is_done": True,
            "actual_deviation": "温度短暂升高",
        }],
    }, record["row_version"])

    proposed_steps = [{
        "title": "原步骤", "instruction": "补充后的操作", "planned_duration_minutes": 20,
        "checkpoint": "检查细胞状态", "risk": "避免污染", "is_done": True,
        "actual_deviation": "温度短暂升高",
    }, {
        "title": "新增检测", "instruction": "完成终点检测", "planned_duration_minutes": 30,
        "checkpoint": "记录原始文件", "risk": "", "is_done": False, "actual_deviation": "",
    }]
    with app.app_context():
        user_id = User.query.filter_by(email="complete-desktop@example.invalid").one().id
        conversation = AIConversation(user_id=user_id, title="步骤建议")
        db.session.add(conversation)
        db.session.flush()
        message = AIMessage(conversation_id=conversation.id, role="assistant", content="步骤建议")
        db.session.add(message)
        db.session.flush()
        change = AIChangeSet(
            message_id=message.id, target_type="lab_record", target_id=record["id"],
            base_row_version=original["row_version"],
            proposal_json=json.dumps({"steps": proposed_steps}, ensure_ascii=False),
            before_json="{}", source_snapshot_json="{}",
        )
        db.session.add(change)
        db.session.commit()
        change_id = change.id

    service.ai_apply({"id": change_id, "accepted_fields": ["steps"]})
    applied = service.get_record({"id": record["id"]})
    assert [step["title"] for step in applied["steps"]] == ["原步骤", "新增检测"]
    assert applied["steps"][0]["is_done"] is True
    assert applied["steps"][0]["actual_deviation"] == "温度短暂升高"

    service.ai_revert({"id": change_id})
    reverted = service.get_record({"id": record["id"]})
    assert len(reverted["steps"]) == 1
    assert reverted["steps"][0]["title"] == "原步骤"
    assert reverted["steps"][0]["actual_deviation"] == "温度短暂升高"


def test_settings_calendar_and_schema_are_desktop_complete(app):
    service = desktop_service(app)
    project = service.create_project({"title": "日历项目"})
    service.settings_save({"section": "workspace", "name": "我的研究工作区", "timezone": "Asia/Shanghai"})
    service.settings_save({"section": "executor", "name": "张研究员", "role": "实验执行"})
    settings = service.settings_get()
    assert settings["workspace"]["name"] == "我的研究工作区"
    assert settings["about"]["author"] == "面壁者"
    assert settings["executors"][0]["name"] == "张研究员"

    event = service.create_calendar_event({
        "title": "课题讨论", "starts_at": "2026-08-14T10:00:00",
        "ends_at": "2026-08-14T11:00:00", "event_type": "meeting",
        "project_id": project["id"],
    })
    events = service.list_calendar({"start": "2026-08-01", "end": "2026-08-31"})
    assert any(item["source_type"] == "event" and item["source_id"] == event["id"] for item in events)

    service.trash_move({"entity_type": "calendar_event", "id": event["id"]})
    events_after_delete = service.list_calendar({"start": "2026-08-01", "end": "2026-08-31"})
    assert not any(item["source_type"] == "event" and item["source_id"] == event["id"] for item in events_after_delete)

    with app.app_context():
        tables = set(inspect(db.engine).get_table_names())
    assert {
        "literature_item", "library_item", "note", "calendar_event", "weekly_report_entry",
        "ai_change_set", "search_document", "activity_event", "file_operation",
    }.issubset(tables)


def test_literature_and_projects_support_recoverable_bulk_removal(app):
    service = desktop_service(app)
    first_project = service.create_project({"title": "项目批量一", "code": "P-01"})
    second_project = service.create_project({"title": "项目批量二", "code": "P-02"})
    changed = service.project_bulk({
        "ids": [first_project["id"], second_project["id"]], "action": "status", "status": "paused",
    })
    assert changed == {"updated": 2, "skipped": 0}
    assert {item["status"] for item in service.list_projects({"search": "P-0"})} == {"paused"}

    literature = service.save_literature({"title": "可删除文献", "authors": "Li"})
    service.trash_move({"entity_type": "literature", "id": literature["id"]})
    assert not any(item["id"] == literature["id"] for item in service.list_literature())
    assert any(item["entity_type"] == "literature" and item["id"] == literature["id"] for item in service.trash_list())
    service.trash_restore({"entity_type": "literature", "id": literature["id"]})
    assert service.get_literature({"id": literature["id"]})["title"] == "可删除文献"

    moved = service.project_bulk({"ids": [first_project["id"]], "action": "trash"})
    assert moved == {"updated": 1, "skipped": 0}
    assert not any(item["id"] == first_project["id"] for item in service.list_projects())


def test_desktop_file_export_and_ppt_workflows_use_selected_paths(app, tmp_path):
    service = desktop_service(app)
    project = service.create_project({"title": "导出项目"})
    record = service.create_record({
        "project_id": project["id"], "title": "导出记录", "experiment_date": "2026-08-11",
    })
    record = service.update_record({
        "id": record["id"], "objective": "验证导出内容", "actual_result": "结果已由用户确认。",
        "conclusion": "导出流程正常。",
    }, record["row_version"])

    source = tmp_path / "result.txt"
    source.write_text("result evidence", encoding="utf-8")
    library_item = service.import_library_item({
        "path": str(source), "display_name": "结果证据.txt", "kind": "raw_data",
        "storage_mode": "external", "project_id": project["id"], "record_id": record["id"],
    })
    assert library_item["storage_mode"] == "external"
    assert service.list_library_items({"record_id": record["id"]})[0]["id"] == library_item["id"]

    json_path = tmp_path / "record.json"
    exported = service.export_record({"id": record["id"], "path": str(json_path)})
    assert exported["size_bytes"] > 0
    assert json.loads(json_path.read_text(encoding="utf-8"))["conclusion"] == "导出流程正常。"

    ppt_path = tmp_path / "weekly.pptx"
    generated = service.generate_weekly_ppt({
        "path": str(ppt_path), "record_ids": [record["id"]], "title": "本周研究进展",
    })
    assert generated["size_bytes"] > 0
    from pptx import Presentation
    assert len(Presentation(ppt_path).slides) == 2


def test_desktop_record_word_and_pdf_exports_follow_editor_schema(app, tmp_path):
    service = desktop_service(app)
    project = service.create_project({"title": "格式兼容项目", "code": "FMT-01"})
    record = service.create_record({
        "project_id": project["id"], "title": "类器官构建记录", "experiment_date": "2026-08-12",
    })
    record = service.update_record({
        "id": record["id"], "objective": "验证构建流程。", "background": "研究背景。",
        "hypothesis": "形成稳定类器官。", "design": "按方案执行。", "materials_conditions": "培养条件。",
        "expected_result": "观察形态。", "actual_process_summary": "完成观察。", "actual_result": "形成类器官。",
        "analysis": "支持假设。", "conclusion": "流程可行。", "next_steps": "开展重复实验。",
    }, record["row_version"])

    docx_path = tmp_path / "record.docx"
    pdf_path = tmp_path / "record.pdf"
    assert service.export_record({"id": record["id"], "path": str(docx_path)})["size_bytes"] > 0
    assert service.export_record({"id": record["id"], "path": str(pdf_path)})["size_bytes"] > 0

    from docx import Document
    document = Document(docx_path)
    docx_text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert all(label in docx_text for label in ("基本信息", "背景、目的与假设", "实验步骤", "结构化参数", "原始数据与附件"))
    assert len(document.tables) >= 8
    assert "R/LAB" not in docx_text

    from pypdf import PdfReader
    pdf_text = "\n".join(page.extract_text() or "" for page in PdfReader(str(pdf_path)).pages)
    assert all(label in pdf_text for label in ("实验记录报告", "背景、目的与假设", "实验步骤", "结构化参数", "人工核验提示"))
    assert "R/LAB" not in pdf_text


def test_weekly_library_imports_managed_versions_and_exports_copies(app, tmp_path):
    service = desktop_service(app)
    project = service.create_project({"title": "周报文件项目"})
    first_source = tmp_path / "week-32-v1.pptx"
    first_source.write_bytes(b"weekly version one")

    first = service.import_weekly_file({
        "path": str(first_source),
        "title": "第 32 周研究进展",
        "report_date": "2026-08-09",
        "period_start": "2026-08-03",
        "period_end": "2026-08-09",
        "project_id": project["id"],
        "storage_mode": "managed",
        "status": "submitted",
        "summary": "完成第一版周总结。",
    })
    assert first["current_file"]["version_number"] == 1
    assert first["current_file"]["storage_mode"] == "managed"
    first_path = service.weekly_file_path({
        "report_id": first["id"], "file_id": first["current_file"]["id"],
    })
    assert first_path["path"] != str(first_source)
    assert Path(first_path["path"]).read_bytes() == b"weekly version one"

    second_source = tmp_path / "week-32-v2.pdf"
    second_source.write_bytes(b"weekly version two")
    second = service.import_weekly_file({
        "report_id": first["id"],
        "path": str(second_source),
        "title": first["title"],
        "report_date": first["report_date"],
        "period_start": first["period_start"],
        "period_end": first["period_end"],
        "project_id": project["id"],
        "storage_mode": "managed",
        "status": "reviewed",
        "summary": "完成第二版并补充批注。",
    }, first["row_version"])
    assert [item["version_number"] for item in second["files"]] == [2, 1]
    assert second["current_file"]["original_name"] == "week-32-v2.pdf"

    target = tmp_path / "saved-copy.pdf"
    exported = service.export_weekly_file({
        "report_id": second["id"],
        "file_id": second["current_file"]["id"],
        "path": str(target),
    })
    assert exported["size_bytes"] == len(b"weekly version two")
    assert target.read_bytes() == b"weekly version two"
    with app.app_context():
        assert WeeklyReport.query.count() == 1
        assert LibraryItem.query.filter_by(kind="report").count() == 2
